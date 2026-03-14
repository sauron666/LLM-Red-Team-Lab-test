from __future__ import annotations
import json
from pathlib import Path
from typing import Any, Dict, List
from datetime import datetime
from jinja2 import Template
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as pdf_canvas

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .db import LabDB
from .visuals import build_heatmap
HTML_TEMPLATE = Template("""
<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>LLM Attack Lab Report — Run {{ run.id }}</title>
<style>
  :root {
    --bg: #09090d; --bg2: #0f0f1a; --bg3: #14141f;
    --border: #1e2040; --accent: #00d4ff; --accent2: #a855f7;
    --text: #c8e0ff; --text2: #7090b0; --text3: #405070;
    --crit: #ff2244; --high: #ff6b35; --med: #ffcc00;
    --low: #00ff88; --info: #00d4ff; --green: #00ff88;
    --font: 'Consolas','Cascadia Code','Courier New',monospace;
    --font-ui: 'Segoe UI','Inter',sans-serif;
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { background: var(--bg); color: var(--text); font-family: var(--font-ui); font-size: 13px; line-height: 1.6; }
  .scanlines { position: fixed; inset: 0; background: repeating-linear-gradient(0deg,transparent,transparent 2px,rgba(0,0,0,0.03) 2px,rgba(0,0,0,0.03) 4px); pointer-events: none; z-index: 9999; }
  header { background: linear-gradient(135deg, #060b12 0%, #0f0f1a 100%); border-bottom: 2px solid var(--accent); padding: 24px 32px; display: flex; align-items: center; gap: 20px; }
  .logo { font-family: var(--font); font-size: 22px; font-weight: 900; color: var(--accent); text-shadow: 0 0 20px rgba(0,212,255,0.5); letter-spacing: 3px; }
  .logo-sub { font-size: 11px; color: var(--text2); letter-spacing: 2px; margin-top: 4px; }
  .badge { padding: 3px 10px; border-radius: 4px; font-family: var(--font); font-size: 10px; font-weight: 700; letter-spacing: 1.5px; border: 1px solid; }
  .badge-crit { background: rgba(255,34,68,.15); color: var(--crit); border-color: rgba(255,34,68,.4); }
  .badge-high { background: rgba(255,107,53,.15); color: var(--high); border-color: rgba(255,107,53,.4); }
  .badge-med  { background: rgba(255,204,0,.15); color: var(--med); border-color: rgba(255,204,0,.4); }
  .badge-low  { background: rgba(0,255,136,.15); color: var(--low); border-color: rgba(0,255,136,.4); }
  .badge-info { background: rgba(0,212,255,.15); color: var(--info); border-color: rgba(0,212,255,.4); }
  .container { max-width: 1400px; margin: 0 auto; padding: 24px 32px; }
  .metrics-grid { display: grid; grid-template-columns: repeat(6,1fr); gap: 12px; margin-bottom: 24px; }
  .metric-card { background: var(--bg2); border: 1px solid var(--border); border-radius: 8px; padding: 16px; position: relative; overflow: hidden; }
  .metric-card::before { content: ''; position: absolute; top: 0; left: 0; right: 0; height: 2px; }
  .metric-card.c-cyan::before { background: var(--accent); }
  .metric-card.c-crit::before { background: var(--crit); }
  .metric-card.c-high::before { background: var(--high); }
  .metric-card.c-med::before  { background: var(--med); }
  .metric-card.c-low::before  { background: var(--low); }
  .metric-card.c-purple::before { background: var(--accent2); }
  .metric-label { font-size: 9px; letter-spacing: 2px; color: var(--text3); text-transform: uppercase; font-family: var(--font); margin-bottom: 8px; }
  .metric-value { font-family: var(--font); font-size: 28px; font-weight: 900; line-height: 1; }
  .metric-card.c-cyan .metric-value { color: var(--accent); }
  .metric-card.c-crit .metric-value { color: var(--crit); }
  .metric-card.c-high .metric-value { color: var(--high); }
  .metric-card.c-med .metric-value  { color: var(--med); }
  .metric-card.c-low .metric-value  { color: var(--low); }
  .metric-card.c-purple .metric-value { color: var(--accent2); }
  .metric-sub { font-size: 10px; color: var(--text3); font-family: var(--font); margin-top: 4px; }
  .section { background: var(--bg2); border: 1px solid var(--border); border-radius: 8px; margin-bottom: 20px; overflow: hidden; }
  .section-header { display: flex; align-items: center; justify-content: space-between; padding: 14px 20px; border-bottom: 1px solid var(--border); background: var(--bg3); }
  .section-title { font-family: var(--font); font-size: 11px; font-weight: 700; letter-spacing: 2px; color: var(--accent); text-transform: uppercase; }
  .section-body { padding: 20px; }
  table { width: 100%; border-collapse: collapse; }
  th { text-align: left; padding: 8px 12px; font-size: 9px; letter-spacing: 2px; color: var(--text3); text-transform: uppercase; font-family: var(--font); border-bottom: 2px solid var(--border); background: var(--bg3); }
  td { padding: 10px 12px; border-bottom: 1px solid rgba(30,32,64,0.6); font-size: 12px; }
  tr:hover td { background: rgba(0,212,255,0.03); }
  tr.sev-critical td:first-child { border-left: 3px solid var(--crit); }
  tr.sev-high td:first-child { border-left: 3px solid var(--high); }
  tr.sev-medium td:first-child { border-left: 3px solid var(--med); }
  tr.sev-low td:first-child { border-left: 3px solid var(--low); }
  .sev { padding: 2px 8px; border-radius: 3px; font-family: var(--font); font-size: 9px; font-weight: 700; letter-spacing: 1px; }
  .sev-critical { background: rgba(255,34,68,.15); color: var(--crit); border: 1px solid rgba(255,34,68,.3); }
  .sev-high     { background: rgba(255,107,53,.15); color: var(--high); border: 1px solid rgba(255,107,53,.3); }
  .sev-medium   { background: rgba(255,204,0,.15); color: var(--med); border: 1px solid rgba(255,204,0,.3); }
  .sev-low      { background: rgba(0,255,136,.15); color: var(--low); border: 1px solid rgba(0,255,136,.3); }
  .sev-info     { background: rgba(0,212,255,.15); color: var(--info); border: 1px solid rgba(0,212,255,.3); }
  .status-vulnerable { color: var(--crit); font-family: var(--font); font-size: 10px; font-weight: 700; }
  .status-blocked { color: var(--green); font-family: var(--font); font-size: 10px; }
  .status-review { color: var(--med); font-family: var(--font); font-size: 10px; }
  pre { background: var(--bg); border: 1px solid var(--border); border-radius: 4px; padding: 12px; font-family: var(--font); font-size: 11px; overflow-x: auto; color: var(--accent); white-space: pre-wrap; word-break: break-all; }
  code { background: var(--bg); color: var(--accent); padding: 1px 6px; border-radius: 3px; font-family: var(--font); font-size: 11px; }
  .finding-card { background: var(--bg); border: 1px solid var(--border); border-radius: 6px; padding: 16px; margin-bottom: 12px; }
  .finding-title { font-size: 14px; font-weight: 600; color: var(--text); margin-bottom: 8px; }
  .finding-meta { display: flex; gap: 12px; flex-wrap: wrap; margin-bottom: 12px; align-items: center; }
  .meta-item { font-size: 11px; color: var(--text2); }
  .meta-key { color: var(--text3); margin-right: 4px; font-family: var(--font); font-size: 9px; letter-spacing: 1px; }
  .recommendation { background: rgba(0,255,136,0.05); border: 1px solid rgba(0,255,136,0.2); border-radius: 4px; padding: 10px 14px; font-size: 12px; color: var(--low); margin-top: 10px; }
  .risk-bar-container { background: rgba(255,255,255,0.05); border-radius: 4px; height: 6px; margin-top: 4px; overflow: hidden; }
  .risk-bar { height: 100%; border-radius: 4px; }
  .zeroday-badge { background: rgba(168,85,247,0.15); border: 1px solid rgba(168,85,247,0.4); color: var(--accent2); padding: 2px 8px; border-radius: 3px; font-size: 9px; font-family: var(--font); letter-spacing: 1px; animation: glow 2s infinite; }
  @keyframes glow { 0%,100% { box-shadow: 0 0 4px rgba(168,85,247,0.4); } 50% { box-shadow: 0 0 12px rgba(168,85,247,0.8); } }
  .divider { height: 1px; background: var(--border); margin: 16px 0; }
  footer { text-align: center; padding: 20px; color: var(--text3); font-family: var(--font); font-size: 10px; letter-spacing: 2px; border-top: 1px solid var(--border); margin-top: 32px; }
</style>
</head>
<body>
<div class="scanlines"></div>

<header>
  <div>
    <div class="logo">⚔️ LLM RED TEAM LAB</div>
    <div class="logo-sub">AI SECURITY RESEARCH PLATFORM — CONFIDENTIAL REPORT</div>
  </div>
  <div style="margin-left:auto;display:flex;gap:10px;align-items:center;">
    <span class="badge {% if summary.severity == 'critical' %}badge-crit{% elif summary.severity == 'high' %}badge-high{% elif summary.severity == 'medium' %}badge-med{% else %}badge-low{% endif %}">
      RISK: {{ summary.overall_risk }}/100 — {{ (summary.severity or 'INFO')|upper }}
    </span>
    <span style="font-family:monospace;font-size:11px;color:#7090b0;">Run #{{ run.id }} | {{ run.started_at[:19] if run.started_at else 'N/A' }}</span>
  </div>
</header>

<div class="container">

  <!-- METRICS -->
  <div class="metrics-grid">
    <div class="metric-card c-cyan">
      <div class="metric-label">Total Findings</div>
      <div class="metric-value">{{ findings|length }}</div>
      <div class="metric-sub">attack cases run</div>
    </div>
    <div class="metric-card c-crit">
      <div class="metric-label">Critical</div>
      <div class="metric-value">{{ findings|selectattr('severity','equalto','critical')|list|length }}</div>
      <div class="metric-sub">immediate action</div>
    </div>
    <div class="metric-card c-high">
      <div class="metric-label">High</div>
      <div class="metric-value">{{ findings|selectattr('severity','equalto','high')|list|length }}</div>
      <div class="metric-sub">review required</div>
    </div>
    <div class="metric-card c-med">
      <div class="metric-label">Medium</div>
      <div class="metric-value">{{ findings|selectattr('severity','equalto','medium')|list|length }}</div>
      <div class="metric-sub">monitor</div>
    </div>
    <div class="metric-card c-low">
      <div class="metric-label">Blocked</div>
      <div class="metric-value">{{ findings|selectattr('status','equalto','blocked')|list|length }}</div>
      <div class="metric-sub">correctly refused</div>
    </div>
    <div class="metric-card c-purple">
      <div class="metric-label">Risk Score</div>
      <div class="metric-value">{{ summary.overall_risk }}</div>
      <div class="metric-sub">overall /100</div>
    </div>
  </div>

  <!-- RUN INFO -->
  <div class="section">
    <div class="section-header">
      <div class="section-title">📋 Run Configuration</div>
    </div>
    <div class="section-body">
      <table style="max-width:700px;">
        <tr><td style="color:#7090b0;width:160px;">Run ID</td><td><code>#{{ run.id }}</code></td></tr>
        <tr><td style="color:#7090b0;">Mode</td><td>{{ run.mode }}</td></tr>
        <tr><td style="color:#7090b0;">Target ID</td><td>{{ run.target_id }}</td></tr>
        <tr><td style="color:#7090b0;">Intensity</td><td>{{ run.intensity }}</td></tr>
        <tr><td style="color:#7090b0;">Safe Mode</td><td>{{ run.safe_mode }}</td></tr>
        <tr><td style="color:#7090b0;">Status</td><td>{{ run.status }}</td></tr>
        <tr><td style="color:#7090b0;">Started</td><td>{{ run.started_at }}</td></tr>
        <tr><td style="color:#7090b0;">Finished</td><td>{{ run.finished_at }}</td></tr>
      </table>
    </div>
  </div>

  <!-- FINDINGS SUMMARY TABLE -->
  <div class="section">
    <div class="section-header">
      <div class="section-title">🎯 Findings Summary</div>
      <span style="font-family:monospace;font-size:10px;color:#7090b0;">{{ findings|length }} total findings</span>
    </div>
    <div class="section-body" style="padding:0;">
      <table>
        <thead>
          <tr>
            <th>#</th><th>Severity</th><th>Attack</th><th>Category</th>
            <th>Status</th><th>Score</th><th>Evidence</th>
          </tr>
        </thead>
        <tbody>
        {% for f in findings %}
          <tr class="sev-{{ f.severity }}">
            <td style="font-family:monospace;color:#405070;">{{ f.id }}</td>
            <td><span class="sev sev-{{ f.severity }}">{{ f.severity|upper }}</span></td>
            <td style="color:#c8e0ff;">{{ f.attack_title }}</td>
            <td style="font-family:monospace;font-size:10px;color:#7090b0;">{{ f.category }}</td>
            <td><span class="status-{{ f.status or 'inconclusive' }}">{{ (f.status or 'inconclusive')|upper }}</span></td>
            <td style="font-family:monospace;color:#00d4ff;">{{ f.risk_score }}</td>
            <td style="font-family:monospace;font-size:10px;color:#7090b0;">{{ f.evidence_type or '' }}</td>
          </tr>
        {% endfor %}
        </tbody>
      </table>
    </div>
  </div>

  <!-- DETAILED FINDINGS -->
  <div class="section">
    <div class="section-header">
      <div class="section-title">🔬 Detailed Findings</div>
    </div>
    <div class="section-body">
    {% for f in findings %}
    {% if f.severity in ['critical','high'] or f.status == 'vulnerable' %}
      <div class="finding-card">
        <div class="finding-title">
          {{ f.attack_title }}
          {% if f.evidence and f.evidence.get('triage', {}).get('zero_day') %}
          <span class="zeroday-badge">ZERO-DAY</span>
          {% endif %}
        </div>
        <div class="finding-meta">
          <span class="sev sev-{{ f.severity }}">{{ f.severity|upper }}</span>
          <span class="meta-item"><span class="meta-key">CATEGORY</span>{{ f.category }}</span>
          <span class="meta-item"><span class="meta-key">STATUS</span><span class="status-{{ f.status or 'inconclusive' }}">{{ (f.status or '?')|upper }}</span></span>
          <span class="meta-item"><span class="meta-key">RISK SCORE</span>{{ f.risk_score }}/100</span>
          <span class="meta-item"><span class="meta-key">CONFIDENCE</span>{{ f.confidence }}/3</span>
          <span class="meta-item"><span class="meta-key">EVIDENCE TYPE</span>{{ f.evidence_type or 'general_signal' }}</span>
        </div>
        <div class="risk-bar-container">
          <div class="risk-bar" style="width:{{ f.risk_score }}%;background:{% if f.risk_score >= 85 %}var(--crit){% elif f.risk_score >= 70 %}var(--high){% elif f.risk_score >= 50 %}var(--med){% else %}var(--low){% endif %};"></div>
        </div>
        {% if f.recommendation %}
        <div class="recommendation">💡 {{ f.recommendation }}</div>
        {% endif %}
        {% if f.evidence %}
        <div class="divider"></div>
        <pre>{{ f.evidence|tojson(indent=2) if f.evidence else '' }}</pre>
        {% endif %}
      </div>
    {% endif %}
    {% endfor %}
    </div>
  </div>

  <!-- HEATMAP -->
  {% if heatmap %}
  <div class="section">
    <div class="section-header">
      <div class="section-title">🌡️ Severity Heatmap by Category</div>
    </div>
    <div class="section-body" style="padding:0;">
      <table>
        <thead>
          <tr>
            <th>Category</th>
            {% for sev in heatmap.severities %}<th>{{ sev|upper }}</th>{% endfor %}
          </tr>
        </thead>
        <tbody>
        {% for cat in heatmap.categories %}
          <tr>
            <td style="font-family:monospace;font-size:11px;">{{ cat }}</td>
            {% for sev in heatmap.severities %}
            {% set count = heatmap.matrix.get(cat, {}).get(sev, 0) %}
            <td style="text-align:center;{% if count > 0 %}background:{% if sev == 'critical' %}rgba(255,34,68,{{ [count * 0.15, 0.6]|min }}){% elif sev == 'high' %}rgba(255,107,53,{{ [count * 0.15, 0.6]|min }}){% elif sev == 'medium' %}rgba(255,204,0,{{ [count * 0.12, 0.5]|min }}){% else %}rgba(0,255,136,{{ [count * 0.10, 0.4]|min }}){% endif %}{% endif %};">
              {% if count > 0 %}<strong>{{ count }}</strong>{% else %}<span style="color:#405070">—</span>{% endif %}
            </td>
            {% endfor %}
          </tr>
        {% endfor %}
        </tbody>
      </table>
    </div>
  </div>
  {% endif %}

</div>

<footer>
  ⚔️ LLM ATTACK LAB — CONFIDENTIAL SECURITY REPORT |
  Generated: {{ now }} | Run #{{ run.id }} |
  For authorized security testing only
</footer>
</body>
</html>
""")

def export_run_report(db: LabDB, run_id: int, out_dir: Path, formats: List[str]) -> Dict[str, str]:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    run = db.get_run(run_id)
    findings = db.list_findings(run_id=run_id)
    summary = run.get("summary", {}) or {}
    heatmap = build_heatmap(findings)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    prefix = out_dir / f"llm_attack_lab_run_{run_id}_{ts}"
    outputs: Dict[str, str] = {}
    if "html" in formats:
        html_path = str(prefix) + ".html"
        from datetime import datetime as _dt
        import json as _json
        _env = HTML_TEMPLATE.environment if hasattr(HTML_TEMPLATE, 'environment') else None
        if _env:
            _env.filters['tojson'] = lambda v, indent=None: _json.dumps(v, indent=indent, ensure_ascii=False, default=str)
        html = HTML_TEMPLATE.render(run=run, findings=findings, summary=summary, heatmap=heatmap, now=_dt.now().strftime('%Y-%m-%d %H:%M UTC'))
        Path(html_path).write_text(html, encoding="utf-8")
        outputs["html"] = html_path
    if "json" in formats:
        json_path = str(prefix) + ".json"
        payload = {"run": run, "summary": summary, "heatmap": heatmap, "findings": findings}
        Path(json_path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        outputs["json"] = json_path
    if "sarif" in formats:
        sarif_path = str(prefix) + ".sarif.json"
        Path(sarif_path).write_text(json.dumps(_build_sarif(run, findings), ensure_ascii=False, indent=2), encoding="utf-8")
        outputs["sarif"] = sarif_path
    if "pdf" in formats:
        pdf_path = str(prefix) + ".pdf"
        _export_pdf(pdf_path, run, findings, summary, heatmap)
        outputs["pdf"] = pdf_path
    if "docx" in formats:
        if DOCX_AVAILABLE:
            docx_path = str(prefix) + ".docx"
            _export_docx(docx_path, run, findings, summary, heatmap)
            outputs["docx"] = docx_path
        else:
            outputs["docx"] = "ERROR: python-docx not installed"
    return outputs
def _export_docx(path: str, run: Dict[str, Any], findings: List[Dict[str, Any]], summary: Dict[str, Any], heatmap: Dict[str, Any]) -> None:
    doc = Document()
    doc.add_heading(f"LLM Attack Lab Report - Run {run['id']}", 0)
    
    doc.add_heading("Run Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Status: ").bold = True
    p.add_run(f"{run['status']}\n")
    p.add_run(f"Mode: ").bold = True
    p.add_run(f"{run['mode']} (Safe Mode: {bool(run['safe_mode'])}, Intensity: {run['intensity']})\n")
    p.add_run(f"Started: ").bold = True
    p.add_run(f"{run['started_at']}\n")
    if run.get('finished_at'):
        p.add_run(f"Finished: ").bold = True
        p.add_run(f"{run['finished_at']}\n")
    p.add_run(f"Overall Risk: ").bold = True
    p.add_run(f"{summary.get('overall_risk', 0)} / {summary.get('severity', 'info')}")
    
    doc.add_heading("Category Summary", level=1)
    for cat, row in (summary.get("by_category") or {}).items():
        doc.add_paragraph(f"• {cat}: {row.get('count')} issues (Max Score: {row.get('max_score')}, Avg: {row.get('avg_score')})")
        
    doc.add_heading("Findings", level=1)
    if not findings:
        doc.add_paragraph("No findings.")
    for f in findings:
        doc.add_heading(f.get('attack_title', 'Unknown Attack'), level=2)
        doc.add_paragraph(f"Category: {f.get('category')} | Severity: {f.get('severity')} | Risk Score: {f.get('risk_score')} | Confidence: {f.get('confidence')}")
        doc.add_paragraph(f"Recommendation: {f.get('recommendation', 'N/A')}")
        excerpt = ((f.get("transcript") or {}).get("response") or "").replace("\n", " ").strip()
        if len(excerpt) > 500: excerpt = excerpt[:500] + "..."
        p_exc = doc.add_paragraph(f"Response Excerpt:\n{excerpt}")
        p_exc.style = 'Intense Quote'
        
    doc.save(path)
def _build_sarif(run: Dict[str, Any], findings: List[Dict[str, Any]]) -> Dict[str, Any]:
    rules = []
    seen_rules = set()
    results = []
    level_map = {"critical": "error", "high": "error", "medium": "warning", "low": "note", "info": "note"}
    for f in findings:
        rule_id = f"LLM.{f.get('category','generic')}.{str(f.get('attack_id','attack')).replace(' ', '_')}"
        if rule_id not in seen_rules:
            rules.append({
                "id": rule_id,
                "name": f.get("attack_title", rule_id),
                "shortDescription": {"text": f.get("category", "LLM finding")},
                "fullDescription": {"text": f.get("recommendation", "LLM Attack Lab finding")},
                "properties": {
                    "category": f.get("category"),
                    "risk_score": f.get("risk_score"),
                },
            })
            seen_rules.add(rule_id)
        transcript = f.get("transcript") or {}
        evidence = f.get("evidence") or {}
        msg = f"{f.get('attack_title')} | risk={f.get('risk_score')} | conf={f.get('confidence')}"
        results.append({
            "ruleId": rule_id,
            "level": level_map.get(str(f.get("severity", "info")).lower(), "warning"),
            "message": {"text": msg},
            "locations": [{
                "physicalLocation": {
                    "artifactLocation": {"uri": f"llm-attack-lab://run/{run['id']}"},
                    "region": {"startLine": 1}
                }
            }],
            "properties": {
                "attack_id": f.get("attack_id"),
                "category": f.get("category"),
                "status": f.get("status", ((f.get("evidence") or {}).get("triage") or {}).get("status")),
                "evidence_type": f.get("evidence_type", ((f.get("evidence") or {}).get("triage") or {}).get("evidence_type")),
                "severity": f.get("severity"),
                "risk_score": f.get("risk_score"),
                "confidence": f.get("confidence"),
                "response_excerpt": str((evidence.get("response_excerpt") or transcript.get("response") or ""))[:1000],
            }
        })
    return {
        "$schema": "https://json.schemastore.org/sarif-2.1.0.json",
        "version": "2.1.0",
        "runs": [{
            "tool": {
                "driver": {
                    "name": "LLM Attack Lab",
                    "version": "0.2",
                    "informationUri": "https://example.local/llm-attack-lab",
                    "rules": rules,
                }
            },
            "automationDetails": {"id": f"run-{run['id']}"},
            "results": results,
        }],
    }
def _export_pdf(path: str, run: Dict[str, Any], findings: List[Dict[str, Any]], summary: Dict[str, Any], heatmap: Dict[str, Any]) -> None:
    c = pdf_canvas.Canvas(path, pagesize=A4)
    width, height = A4
    y = height - 40
    def line(text: str, step: int = 14):
        nonlocal y
        c.drawString(40, y, text[:120])
        y -= step
        if y < 60:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 40
    c.setFont("Helvetica-Bold", 16)
    line("LLM Attack Lab Report", 22)
    c.setFont("Helvetica", 10)
    line(f"Run ID: {run['id']}")
    line(f"Status: {run['status']} | Mode: {run['mode']} | Safe Mode: {bool(run['safe_mode'])} | Intensity: {run['intensity']}")
    line(f"Started: {run['started_at']} | Finished: {run.get('finished_at')}")
    line(f"Overall Risk: {summary.get('overall_risk', 0)} ({summary.get('severity', 'info')})", 18)
    c.setFont("Helvetica-Bold", 12)
    line("Category Summary", 18)
    c.setFont("Helvetica", 10)
    for cat, row in (summary.get("by_category") or {}).items():
        line(f"- {cat}: count={row.get('count')} max={row.get('max_score')} avg={row.get('avg_score')}")
    c.setFont("Helvetica-Bold", 12)
    line("Heatmap", 18)
    c.setFont("Helvetica", 10)
    line("Category | " + " | ".join(heatmap.get("severities", [])))
    for cat in heatmap.get("categories", []):
        row = heatmap["matrix"].get(cat, {})
        line(f"{cat} | " + " | ".join(str(row.get(sev, 0)) for sev in heatmap.get("severities", [])))
    c.setFont("Helvetica-Bold", 12)
    line("Findings", 18)
    c.setFont("Helvetica", 10)
    if not findings:
        line("No findings.")
    else:
        for f in findings:
            line(f"[{str(f.get('status','vulnerable')).upper()} / {str(f['severity']).upper()}] {f['attack_title']} ({f['category']}) score={f['risk_score']} conf={f['confidence']}")
            reco = f.get("recommendation", "")
            line(f"  Recommendation: {reco}", 12)
            excerpt = ((f.get("transcript") or {}).get("response") or "").replace("\n", " ")
            line(f"  Response excerpt: {excerpt[:95]}", 12)
            y -= 4
    c.save()